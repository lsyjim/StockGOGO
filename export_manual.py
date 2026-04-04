"""
export_manual.py
產生「投資評級操作手冊」的 PDF 與 Word 版本
專業投資報告排版風格
"""

import os

# ─── 顏色與字體常量 ─────────────────────────────────────────────────────────

NAVY      = (0.09, 0.18, 0.31)      # 深藍（主標題底色）
DARK_GRAY = (0.18, 0.18, 0.18)      # 深灰（正文）
MID_GRAY  = (0.45, 0.45, 0.45)      # 中灰（副標題）
GOLD      = (0.76, 0.60, 0.19)      # 金色（強調）
TABLE_HDR = (0.13, 0.26, 0.44)      # 表頭藍
TABLE_ALT = (0.93, 0.95, 0.98)      # 交替行淡藍
WHITE     = (1.0,  1.0,  1.0)
RED_DARK  = (0.65, 0.10, 0.10)

FONT_PATH = "/System/Library/Fonts/STHeiti Medium.ttc"


# ══════════════════════════════════════════════════════════════════════════════
#  PDF 版本（reportlab）
# ══════════════════════════════════════════════════════════════════════════════

def build_pdf(output_path: str):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, PageBreak, KeepTogether
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus.flowables import Flowable

    # 註冊中文字體
    pdfmetrics.registerFont(TTFont("STHeiti", FONT_PATH))
    pdfmetrics.registerFont(TTFont("STHeiti-B", FONT_PATH))  # 粗體用同字體

    W, H = A4
    MARGIN = 20 * mm

    # ── 樣式定義 ────────────────────────────────────────────────────────────

    def _c(r, g, b):
        return colors.Color(r, g, b)

    s_cover_title = ParagraphStyle("cover_title",
        fontName="STHeiti", fontSize=28, leading=36,
        textColor=WHITE, alignment=1, spaceAfter=6)

    s_cover_sub = ParagraphStyle("cover_sub",
        fontName="STHeiti", fontSize=13, leading=18,
        textColor=_c(*GOLD), alignment=1, spaceAfter=4)

    s_cover_meta = ParagraphStyle("cover_meta",
        fontName="STHeiti", fontSize=10, leading=14,
        textColor=_c(0.75, 0.75, 0.75), alignment=1)

    s_h1 = ParagraphStyle("h1",
        fontName="STHeiti", fontSize=15, leading=20,
        textColor=_c(*NAVY), spaceBefore=14, spaceAfter=6,
        borderPad=4)

    s_h2 = ParagraphStyle("h2",
        fontName="STHeiti", fontSize=12, leading=16,
        textColor=WHITE, spaceBefore=10, spaceAfter=4,
        backColor=_c(*TABLE_HDR), borderPad=5, leftIndent=0)

    s_h3 = ParagraphStyle("h3",
        fontName="STHeiti", fontSize=11, leading=15,
        textColor=_c(*GOLD), spaceBefore=8, spaceAfter=3,
        leftIndent=8)

    s_body = ParagraphStyle("body",
        fontName="STHeiti", fontSize=9.5, leading=16,
        textColor=_c(*DARK_GRAY), spaceAfter=4, leftIndent=4)

    s_bullet = ParagraphStyle("bullet",
        fontName="STHeiti", fontSize=9.5, leading=15,
        textColor=_c(*DARK_GRAY), leftIndent=16, spaceAfter=2,
        bulletIndent=6)

    s_note = ParagraphStyle("note",
        fontName="STHeiti", fontSize=8.5, leading=13,
        textColor=_c(*MID_GRAY), leftIndent=8, spaceAfter=3,
        italic=True)

    s_footer = ParagraphStyle("footer",
        fontName="STHeiti", fontSize=7.5, leading=11,
        textColor=_c(*MID_GRAY), alignment=1)

    s_tbl_hdr = ParagraphStyle("tbl_hdr",
        fontName="STHeiti", fontSize=9, leading=12,
        textColor=WHITE, alignment=1)

    s_tbl_cell = ParagraphStyle("tbl_cell",
        fontName="STHeiti", fontSize=8.5, leading=13,
        textColor=_c(*DARK_GRAY))

    s_tbl_cell_c = ParagraphStyle("tbl_cell_c",
        fontName="STHeiti", fontSize=8.5, leading=13,
        textColor=_c(*DARK_GRAY), alignment=1)

    # ── 封面 Flowable ────────────────────────────────────────────────────────

    class CoverBlock(Flowable):
        def __init__(self, w, h):
            super().__init__()
            self.w = w
            self.h = h

        def wrap(self, *args):
            return self.w, self.h

        def draw(self):
            c = self.canv
            # 深藍底
            c.setFillColorRGB(*NAVY)
            c.rect(0, 0, self.w, self.h, fill=1, stroke=0)
            # 金線裝飾（頂部）
            c.setStrokeColorRGB(*GOLD)
            c.setLineWidth(2.5)
            c.line(30, self.h - 18, self.w - 30, self.h - 18)
            # 金線裝飾（底部）
            c.line(30, 18, self.w - 30, 18)
            # 主標題
            c.setFillColorRGB(*WHITE)
            c.setFont("STHeiti", 30)
            c.drawCentredString(self.w / 2, self.h - 70, "投資評級操作手冊")
            # 英文副標
            c.setFillColorRGB(*GOLD)
            c.setFont("STHeiti", 13)
            c.drawCentredString(self.w / 2, self.h - 96,
                                "Investment Grade Operating Manual")
            # 分隔線
            c.setStrokeColorRGB(*GOLD)
            c.setLineWidth(0.8)
            c.line(self.w / 2 - 80, self.h - 110, self.w / 2 + 80, self.h - 110)
            # 系統名稱
            c.setFillColorRGB(0.75, 0.75, 0.75)
            c.setFont("STHeiti", 10)
            c.drawCentredString(self.w / 2, self.h - 128,
                                "StockGOGO V2  ·  三層決策引擎")
            # 版本 / 日期
            c.setFont("STHeiti", 9)
            c.drawCentredString(self.w / 2, 36, "Version 2.0  |  2026-04")

    # ── 表格工具 ─────────────────────────────────────────────────────────────

    def make_table(headers, rows, col_widths, centered_cols=None):
        centered_cols = centered_cols or []
        hdr_row = [Paragraph(h, s_tbl_hdr) for h in headers]
        data = [hdr_row]
        for i, row in enumerate(rows):
            cells = []
            for j, cell in enumerate(row):
                st = s_tbl_cell_c if j in centered_cols else s_tbl_cell
                cells.append(Paragraph(str(cell), st))
            data.append(cells)

        tbl = Table(data, colWidths=col_widths, repeatRows=1)
        style_cmds = [
            ("BACKGROUND",  (0, 0), (-1, 0),  _c(*TABLE_HDR)),
            ("TEXTCOLOR",   (0, 0), (-1, 0),  colors.white),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, _c(*TABLE_ALT)]),
            ("GRID",        (0, 0), (-1, -1),  0.4, _c(0.80, 0.83, 0.88)),
            ("FONTNAME",    (0, 0), (-1, -1),  "STHeiti"),
            ("FONTSIZE",    (0, 0), (-1, -1),  8.5),
            ("TOPPADDING",  (0, 0), (-1, -1),  5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1),  6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("VALIGN",      (0, 0), (-1, -1),  "MIDDLE"),
        ]
        tbl.setStyle(TableStyle(style_cmds))
        return tbl

    def hr():
        return HRFlowable(width="100%", thickness=0.5,
                          color=_c(0.75, 0.78, 0.85), spaceAfter=4, spaceBefore=2)

    def sp(h=4):
        return Spacer(1, h * mm)

    def h1(text):
        return Paragraph(f"<b>{text}</b>", s_h1)

    def h2(text):
        return Paragraph(f"  {text}", s_h2)

    def h3(text):
        return Paragraph(f"▌ {text}", s_h3)

    def body(text):
        return Paragraph(text, s_body)

    def bullet(text):
        return Paragraph(f"• {text}", s_bullet)

    def note(text):
        return Paragraph(f"※ {text}", s_note)

    # ── 頁首頁尾 ────────────────────────────────────────────────────────────

    def on_page(canvas, doc):
        canvas.saveState()
        # 頁首
        canvas.setFillColorRGB(*NAVY)
        canvas.rect(0, H - 10 * mm, W, 10 * mm, fill=1, stroke=0)
        canvas.setFillColorRGB(*WHITE)
        canvas.setFont("STHeiti", 8)
        canvas.drawString(MARGIN, H - 6.5 * mm, "StockGOGO V2  投資評級操作手冊")
        canvas.drawRightString(W - MARGIN, H - 6.5 * mm, "Investment Grade Manual")
        # 頁尾
        canvas.setFillColorRGB(*MID_GRAY)
        canvas.setFont("STHeiti", 7.5)
        canvas.drawString(MARGIN, 8 * mm, "機密文件・僅供內部使用・版本 v2.0")
        canvas.drawRightString(W - MARGIN, 8 * mm,
                               f"第 {doc.page} 頁")
        canvas.setStrokeColorRGB(*GOLD)
        canvas.setLineWidth(0.6)
        canvas.line(MARGIN, 11 * mm, W - MARGIN, 11 * mm)
        canvas.restoreState()

    # ── 組裝內容 ────────────────────────────────────────────────────────────

    story = []

    # 封面（獨立頁，無頁首頁尾）
    story.append(CoverBlock(W - 2 * MARGIN, 220 * mm))
    story.append(PageBreak())

    # ── 一、架構概覽 ──────────────────────────────────────────────────────
    story += [
        h1("一、評級架構概覽"),
        hr(),
        body("StockGOGO V2 採用「三層過濾 → 評級輸出」架構，每層均設有否決門檻，"
             "不通過則直接輸出否決結果，不進入下一層運算。此設計確保分析邏輯清晰、"
             "各層因子不互相干擾。"),
        sp(3),
    ]

    # 流程圖（文字版）
    flow_data = [
        ["", "Layer 1", "方向分", "均線排列是否偏多？", "< 40 → SKIP"],
        ["", "Layer 2", "位置分", "現在位置是否合理？", "< 40 → WAIT"],
        ["", "Layer 3", "時機分", "今天是否有進場觸發？", "→ A / B / C / X"],
        ["", "賣訊", "優先", "三盤跌破 / 形態頭部 / 籌碼惡化", "→ SELL / HOLD"],
    ]
    story.append(make_table(
        ["", "層級", "評分維度", "核心問題", "輸出"],
        flow_data,
        [8 * mm, 18 * mm, 18 * mm, 70 * mm, 42 * mm],
        centered_cols=[0, 1, 2]
    ))
    story.append(sp(4))

    # 評級總覽表
    story.append(h3("評級總覽"))
    overview = [
        ["A 主攻",  "STRONG_BUY", "Layer 3", "70–100", "立即進場，50–80% 部位"],
        ["B 追蹤",  "BUY",        "Layer 3", "55–69",  "輕倉試單，等確認再加碼"],
        ["C 觀察",  "HOLD",       "Layer 3", "40–54",  "列入名單，目前不開倉"],
        ["X 無訊號","HOLD",       "Layer 3", "40–54",  "同 C 級，不操作"],
        ["WAIT",    "WAIT",       "Layer 2 否決", "30–49", "趨勢對但位置過熱，等拉回"],
        ["SKIP",    "SKIP",       "Layer 1 否決", "0–39",  "趨勢空頭，完全跳過"],
        ["SELL 防守","SELL",      "賣訊優先", "—",    "三盤跌破，次日開盤全出"],
        ["SELL 停利","HOLD/SELL", "賣訊優先", "—",    "過熱賣訊，分批獲利了結"],
        ["SELL 反轉","SELL",      "賣訊優先", "—",    "頭部確立或法人連賣 7 天"],
    ]
    story.append(make_table(
        ["評級", "動作代碼", "觸發層", "分數區間", "動作摘要"],
        overview,
        [22 * mm, 26 * mm, 24 * mm, 20 * mm, 62 * mm],
        centered_cols=[1, 2, 3]
    ))
    story.append(PageBreak())

    # ── 二、買進評級 ──────────────────────────────────────────────────────
    story += [h1("二、買進評級詳細說明"), hr()]

    # A 級
    story += [
        h2("A 級  主攻（立即進場）"),
        sp(2),
        body("<b>定義：</b>方向分通過 + 位置分通過 + 今天有明確的量價觸發訊號。"
             "三層全部到位，是最強力的進場訊號。"),
        sp(2),
        h3("觸發條件（任一即可）"),
    ]
    for t in ["三盤突破 + 帶量確認（收盤突破近三日高點，成交量明顯放大）",
              "VP05 帶量突破訊號（量價分析系統確認）",
              "底部形態（W底 / 頭肩底）突破頸線，且為剛確立狀態"]:
        story.append(bullet(t))

    story += [sp(2), h3("操作動作")]
    a_tbl = [
        ["進場時機", "當日或次日開盤，不等待"],
        ["部位大小", "50–80% 標準部位"],
        ["停損設定", "進場前低點，或 MA20 下方 2–3%"],
        ["停利目標", "以 RR ≥ 2 為基準，參考形態測幅"],
        ["注意事項", "當日已漲 >5%，降為 B 級，等回檔再進"],
    ]
    story.append(make_table(["項目", "建議"], a_tbl,
                            [40 * mm, 114 * mm]))
    story += [
        sp(2),
        note("A 級不是「感覺不錯」，而是量價到位。方向對但沒有觸發訊號，"
             "最多只是 B 級。"),
        sp(4),
    ]

    # B 級
    story += [
        h2("B 級  追蹤（等待確認）"),
        sp(2),
        body("<b>定義：</b>方向分通過 + 位置分通過 + 訊號形成但尚未完整確認。"
             "值得買，但今天不是最佳時機。"),
        sp(2),
        h3("觸發條件（任一即可）"),
    ]
    for t in ["三盤突破但量能不足（無帶量確認）",
              "超跌反彈左側訊號（RSI 超賣區出現止跌量）",
              "底部形態形成中，且距離頸線 < 3%"]:
        story.append(bullet(t))

    story += [sp(2), h3("操作動作")]
    b_tbl = [
        ["進場時機", "等待：① 帶量突破確認 ② 次日跳空開高 ③ 量縮回踩不破支撐"],
        ["部位大小", "30–50% 輕倉試單，升格 A 後再加碼"],
        ["停損設定", "比 A 級稍寬，設在近期低點或 MA60"],
        ["停利目標", "RR ≥ 1.5，目標相對保守"],
        ["加碼條件", "出現 A 級訊號後，才加碼到滿倉"],
    ]
    story.append(make_table(["項目", "建議"], b_tbl,
                            [40 * mm, 114 * mm]))
    story += [
        sp(2),
        note("B 級不是「馬上買」，是「值得盯著，等訊號」。"
             "停損設太緊容易被洗，建議略寬。"),
        sp(4),
    ]

    # C / X
    story += [
        h2("C 級  觀察 ／ X 無訊號（不操作）"),
        sp(2),
        body("<b>定義：</b>方向分與位置分均通過，但時機層沒有具體觸發條件。"
             "C 為多頭環境但無催化，X 為連多頭環境都不明確。"),
        sp(2),
        h3("操作動作"),
    ]
    cx_tbl = [
        ["現在做什麼", "加入自選股，設定價格提醒"],
        ["升格條件",   "帶量突破 → 升 A；左側訊號或接近頸線 → 升 B"],
        ["已持有倉位", "無賣訊則繼續持有，但不加碼"],
        ["X vs C",     "X 比 C 更不確定，升格機率較低，優先序較後"],
    ]
    story.append(make_table(["項目", "說明"], cx_tbl,
                            [40 * mm, 114 * mm]))
    story.append(PageBreak())

    # ── 三、過濾評級 ──────────────────────────────────────────────────────
    story += [h1("三、過濾評級詳細說明"), hr()]

    # WAIT
    story += [
        h2("WAIT  等待拉回（Layer 2 否決）"),
        sp(2),
        body("<b>定義：</b>方向分通過（趨勢向上），但位置分 < 40（位置過熱或 RR 不足）。"
             "系統幫你踩煞車，避免追高。"),
        sp(2),
        h3("常見觸發情況"),
    ]
    for t in ["乖離率 > +10%（距 MA20 超漲 10%）",
              "RSI > 80（技術面嚴重過熱）",
              "風險報酬比 < 1.5（上方空間不足）",
              "乖離率 > +15%（天花板啟動，位置分硬封頂 ≤ 20）"]:
        story.append(bullet(t))

    story += [sp(2), h3("操作動作")]
    w_tbl = [
        ["現在做什麼", "絕對不追，耐心等待"],
        ["等什麼",     "乖離回歸、RSI 降溫、或出現回檔量縮"],
        ["升格條件",   "乖離回到 -5%~+3%，RSI 回到 40–65，升為 B 或 A"],
        ["最常犯錯誤", "「趨勢很好應該繼續漲」—— 此理由在 WAIT 時完全無效"],
    ]
    story.append(make_table(["項目", "說明"], w_tbl,
                            [40 * mm, 114 * mm]))
    story += [
        sp(2),
        note("WAIT 讓人痛苦，因為眼睜睜看著繼續漲。"
             "但追高後套牢的損失，遠大於沒買到的機會成本。"),
        sp(4),
    ]

    # SKIP
    story += [
        h2("SKIP  跳過（Layer 1 否決）"),
        sp(2),
        body("<b>定義：</b>均線方向分 < 40，趨勢反對你，任何買進都是逆勢操作。"
             "不是這檔不好，而是現在的趨勢對你不利。"),
        sp(2),
        h3("常見觸發情況"),
    ]
    for t in ["均線完全空頭排列（MA20 < MA60 < MA120 < MA240）",
              "現價在年線以下",
              "4 條均線中，多頭方向 ≤ 1 條"]:
        story.append(bullet(t))

    story += [sp(2), h3("操作動作")]
    sk_tbl = [
        ["現在做什麼", "完全不看這檔，從關注名單移除"],
        ["重新關注條件", "方向分升到 40 以上（至少 2/4 均線轉多頭）"],
        ["已持有倉位", "SKIP 本身不是賣訊，需同時出現賣訊評級才出場"],
        ["最常犯錯誤", "「已經跌很多應該反彈」—— SKIP 時此理由無效"],
    ]
    story.append(make_table(["項目", "說明"], sk_tbl,
                            [40 * mm, 114 * mm]))

    story += [sp(3), h3("SKIP vs. WAIT 差異比較")]
    diff_tbl = [
        ["問題所在", "趨勢空頭或混亂",    "位置過熱，追高風險"],
        ["趨勢狀態", "空頭確認",           "多頭確認"],
        ["解決辦法", "等趨勢轉向（數週）", "等位置拉回（數天）"],
        ["緊急程度", "完全退場觀察",       "耐心等待位置"],
    ]
    story.append(make_table(
        ["比較項目", "SKIP", "WAIT"],
        diff_tbl,
        [40 * mm, 58 * mm, 58 * mm],
        centered_cols=[1, 2]
    ))
    story.append(PageBreak())

    # ── 四、賣出評級 ──────────────────────────────────────────────────────
    story += [h1("四、賣出評級詳細說明"), hr(),
              body("賣訊優先於所有買訊。三種賣出類型各有明確觸發條件與執行節奏，"
                   "不混用、不猜測。"),
              sp(3)]

    # 防守型
    story += [
        h2("SELL 防守型  立即出場"),
        sp(2),
        body("<b>定義：</b>已有倉位，現在繼續持有的風險已超過潛在獲利。"
             "最高優先級，代表趨勢已轉向或停損被觸發，須立即執行。"),
        sp(2),
        h3("觸發條件（任一即可）"),
    ]
    for t in ["三盤跌破（收盤跌破近三日低點，趨勢轉空）",
              "跌破 MA20 超過 2%（收盤在 MA20 × 0.98 以下）"]:
        story.append(bullet(t))

    story += [sp(2), h3("操作動作")]
    def1_tbl = [
        ["執行速度", "收到訊號後，次日開盤確認後即出"],
        ["部位處理", "全部出場，不留半倉「賭反彈」"],
        ["出場方式", "掛市價或略低於市價的限價，確保成交"],
        ["之後動作", "等系統重新評為 B 或 A，再考慮重新進場"],
        ["最常犯錯誤", "「跌破了但明天會反彈」—— 嚴格執行停損是必要紀律"],
    ]
    story.append(make_table(["項目", "建議"], def1_tbl,
                            [40 * mm, 114 * mm]))
    story += [sp(2),
              note("防守型賣訊出現時，繼續持有的期望值已為負。"
                   "停損後漲回的痛苦，遠小於沒停損持續虧損的實際損失。"),
              sp(4)]

    # 停利型
    story += [
        h2("SELL 停利型  獲利了結"),
        sp(2),
        body("<b>定義：</b>已有獲利，技術面顯示上漲動能衰竭。"
             "不是緊急出場，而是主動管理倉位的正向動作。"),
        sp(2),
        h3("觸發條件（任一即可）"),
    ]
    for t in ["過熱賣訊：乖離率 > +15%，且出現長上影線或爆量不漲",
              "移動停利觸發：從最高點回落 5%，或跌破 5 日線"]:
        story.append(bullet(t))

    story += [sp(2), h3("操作動作")]
    pt_tbl = [
        ["執行方式", "不必急，可分批出場（先減 50%，剩餘跟移動停利）"],
        ["是否全出", "視個股強弱，強勢股可留 20–30% 看後續"],
        ["出場標準", "以計畫中的目標價為基準，達標就出"],
        ["最常犯錯誤", "「還會繼續漲」而遲遲不出，最後停利變停損"],
    ]
    story.append(make_table(["項目", "建議"], pt_tbl,
                            [40 * mm, 114 * mm]))
    story.append(sp(4))

    # 反轉型
    story += [
        h2("SELL 反轉型  趨勢反轉出場"),
        sp(2),
        body("<b>定義：</b>整體趨勢已轉向，不只是短線回調，而是結構性改變。"
             "代表多頭週期可能結束，需重新評估持有理由。"),
        sp(2),
        h3("觸發條件（任一即可）"),
    ]
    for t in ["頭部形態確立（M頭 / 頭肩頂跌破頸線確認）",
              "法人連續賣超 ≥ 7 天（籌碼嚴重惡化）"]:
        story.append(bullet(t))

    story += [sp(2), h3("操作動作")]
    rev_tbl = [
        ["執行方式", "視情況分批或全部出場"],
        ["頭部形態觸發", "跌破頸線當天或次日，盡速出場"],
        ["籌碼惡化觸發", "觀察 2–3 天，趨勢未改善則出場"],
        ["重新進場條件", "等評級重新回到 B 或 A，且趨勢明確轉多"],
        ["最常犯錯誤", "「基本面沒變壞」—— 技術面反轉先跑，等基本面確認往往太慢"],
    ]
    story.append(make_table(["項目", "建議"], rev_tbl,
                            [40 * mm, 114 * mm]))

    story += [sp(3), h3("三種賣訊比較")]
    sell_cmp = [
        ["防守型", "SELL（urgent）", "三盤跌破 / 跌破 MA20", "高",   "全部出場"],
        ["停利型", "HOLD / SELL",    "過熱賣訊 / 移動停利",  "中",   "分批出場，可留倉"],
        ["反轉型", "SELL（urgent）", "頭部確立 / 法人連賣",  "中高", "盡速出場"],
    ]
    story.append(make_table(
        ["類型", "動作代碼", "觸發條件", "緊急程度", "執行方式"],
        sell_cmp,
        [22 * mm, 26 * mm, 52 * mm, 18 * mm, 38 * mm],
        centered_cols=[1, 3]
    ))
    story.append(PageBreak())

    # ── 五、倉位管理 ──────────────────────────────────────────────────────
    story += [h1("五、倉位管理與籌碼輔助"), hr()]

    story += [h3("評級對應的倉位建議")]
    pos_tbl = [
        ["A 主攻",   "50–80%", "三層全過 + 強觸發，積極進場"],
        ["B 追蹤",   "30–50%", "等確認再加碼，先輕倉"],
        ["C / X",    "0%",     "不開倉，列入追蹤"],
        ["WAIT",     "0%",     "不追高，等位置改善"],
        ["SKIP",     "0%",     "完全不看，從名單移除"],
        ["賣訊出現", "0%",     "依類型執行出場"],
    ]
    story.append(make_table(
        ["評級", "最大持倉", "說明"],
        pos_tbl,
        [25 * mm, 22 * mm, 109 * mm],
        centered_cols=[1]
    ))

    story += [sp(4), h3("加碼邏輯（唯一合理路徑）")]
    add_tbl = [
        ["步驟 1", "B 級出現", "輕倉進場 30–50%"],
        ["步驟 2", "升格為 A 級", "加碼至 50–80%"],
        ["步驟 3", "獲利達 +10%", "啟動移動停利追蹤"],
        ["步驟 4", "停利型賣訊出現", "分批出場"],
    ]
    story.append(make_table(
        ["步驟", "條件", "動作"],
        add_tbl,
        [18 * mm, 50 * mm, 88 * mm]
    ))

    story += [sp(2),
              note("評級下降（A→B→C）時，不加碼，考慮減碼。"
                   "出現任何 SELL 訊號時，只考慮出場，不加碼。"),
              sp(4)]

    story += [h3("籌碼的輔助作用")]
    chip_tbl = [
        ["法人連買 ≥ 5 天", "+2", "同評級中優先選擇"],
        ["法人連買 3–4 天", "+1", "次優先"],
        ["中性",            " 0", "正常排序"],
        ["法人連賣 3–4 天", "−1", "同評級中後選"],
        ["法人連賣 ≥ 5 天", "−2", "B/C 級自動降一級"],
    ]
    story.append(make_table(
        ["籌碼狀態", "排序加成", "實際效果"],
        chip_tbl,
        [60 * mm, 22 * mm, 74 * mm],
        centered_cols=[1]
    ))
    story += [sp(2),
              note("籌碼不決定買賣，只做排序與過濾。"
                   "再好的籌碼，趨勢空頭時仍為 SKIP；再差的籌碼，A 級也不會降為 SKIP。"),
              sp(4)]

    # ── 六、情境對照 ──────────────────────────────────────────────────────
    story += [h1("六、常見情境對照"), hr()]

    cases = [
        ["台積電均線多頭，三盤突破帶量", "A 主攻", "進場 50–80%"],
        ["某股趨勢好，但 RSI=82、乖離+12%", "WAIT", "不追，等拉回"],
        ["某股剛從低點反彈，W底接近頸線", "B 追蹤", "輕倉試單，等突破"],
        ["某股年線以下，均線死亡排列", "SKIP", "完全不看"],
        ["持有中股票出現三盤跌破", "SELL 防守", "次日開盤出場"],
        ["持有中股票漲15%後爆量長上影線", "SELL 停利", "分批出場"],
        ["持有中股票出現 M 頭跌破頸線", "SELL 反轉", "盡速出場"],
        ["持倉評級從 A 降到 C", "C 觀察", "不加碼，持有觀察"],
        ["同時有 3 檔 A 級，資金有限", "A 主攻", "按籌碼 ranking_boost 排序選擇"],
        ["評級 B + 法人連賣 6 天", "C 觀察", "自動降級，降低部位優先序"],
    ]
    story.append(make_table(
        ["情境", "評級", "動作"],
        cases,
        [82 * mm, 24 * mm, 50 * mm],
        centered_cols=[1]
    ))

    story += [sp(4), h1("七、評級升降格含義"), hr()]
    updown = [
        ["升格", "SKIP → WAIT", "趨勢開始轉多，但尚不值得追"],
        ["升格", "WAIT → B/C",  "位置回到合理範圍，進入觀察"],
        ["升格", "C/B → A",     "量價確認，可以動手"],
        ["降格", "A → B",       "觸發訊號消失，可能假突破，考慮減倉"],
        ["降格", "B → C",       "環境轉差或訊號失效，持有但不加碼"],
        ["降格", "C → WAIT/SKIP", "趨勢或位置惡化，已持有者考慮減碼"],
    ]
    story.append(make_table(
        ["方向", "路徑", "含義"],
        updown,
        [18 * mm, 38 * mm, 100 * mm],
        centered_cols=[0]
    ))
    story += [sp(2),
              note("降格不等於馬上賣出。降格是警示，賣出需同時出現 SELL 賣訊。"
                   "沒有 SELL 賣訊的降格，代表持有風險上升，但尚未到出場時機。"),
              sp(6)]

    # 版權頁
    story += [
        HRFlowable(width="100%", thickness=1.0, color=_c(*NAVY)),
        sp(3),
        Paragraph("本手冊為 StockGOGO V2 三層決策引擎配套說明文件", s_footer),
        Paragraph("Version 2.0  ·  2026-04  ·  機密文件，僅供內部使用", s_footer),
    ]

    # ── 輸出 ────────────────────────────────────────────────────────────────
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=14 * mm, bottomMargin=16 * mm,
        title="投資評級操作手冊",
        author="StockGOGO V2",
    )
    doc.build(story, onFirstPage=lambda c, d: None, onLaterPages=on_page)
    print(f"PDF 已輸出：{output_path}")


# ══════════════════════════════════════════════════════════════════════════════
#  Word 版本（python-docx）
# ══════════════════════════════════════════════════════════════════════════════

def build_word(output_path: str):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    doc = Document()

    # ── 頁面設定 ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin  = Cm(2.0)
    section.bottom_margin = Cm(1.8)

    # ── 顏色工具 ──────────────────────────────────────────────────────────
    def rgb(r_f, g_f, b_f):
        return RGBColor(int(r_f * 255), int(g_f * 255), int(b_f * 255))

    C_NAVY  = rgb(*NAVY)
    C_GOLD  = rgb(*GOLD)
    C_GRAY  = rgb(*MID_GRAY)
    C_WHITE = rgb(1, 1, 1)
    C_DARK  = rgb(*DARK_GRAY)
    C_THDR  = rgb(*TABLE_HDR)
    C_TALT  = rgb(*TABLE_ALT)

    # ── 段落工具 ──────────────────────────────────────────────────────────
    def set_para_shading(para, fill_hex):
        """設定段落底色"""
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill_hex)
        pPr.append(shd)

    def set_cell_color(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def hex_color(r_f, g_f, b_f):
        return "{:02X}{:02X}{:02X}".format(
            int(r_f * 255), int(g_f * 255), int(b_f * 255))

    NAV_HEX  = hex_color(*NAVY)
    GOLD_HEX = hex_color(*GOLD)
    THDR_HEX = hex_color(*TABLE_HDR)
    TALT_HEX = hex_color(*TABLE_ALT)

    def add_run(para, text, bold=False, color=None, size=10, italic=False):
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "STHeiti"
        run._r.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
        if color:
            run.font.color.rgb = color
        return run

    def add_heading1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        # 藍色底色
        set_para_shading(p, NAV_HEX)
        run = p.add_run(f"  {text}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "STHeiti"
        run._r.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
        run.font.color.rgb = C_WHITE
        p.paragraph_format.left_indent = Cm(0)

    def add_heading2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
        set_para_shading(p, THDR_HEX)
        run = p.add_run(f"  {text}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = "STHeiti"
        run._r.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
        run.font.color.rgb = C_WHITE

    def add_heading3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.3)
        run = p.add_run(f"▌ {text}")
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = "STHeiti"
        run._r.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
        run.font.color.rgb = C_GOLD

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.2)
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True
            r.font.size = Pt(9.5)
            r.font.name = "STHeiti"
            r._r.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
            r.font.color.rgb = C_DARK
        r2 = p.add_run(text)
        r2.font.size = Pt(9.5)
        r2.font.name = "STHeiti"
        r2._r.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
        r2.font.color.rgb = C_DARK

    def add_bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.name = "STHeiti"
        run._r.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
        run.font.color.rgb = C_DARK

    def add_note(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run(f"※ {text}")
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.name = "STHeiti"
        run._r.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
        run.font.color.rgb = C_GRAY

    def add_table(headers, rows, col_widths_cm, alt=True):
        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"

        # 表頭
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            set_cell_color(hdr_cells[i], THDR_HEX)
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "STHeiti"
            run._r.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
            run.font.color.rgb = C_WHITE
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 資料行
        for ri, row in enumerate(rows):
            cells = tbl.rows[ri + 1].cells
            fill = TALT_HEX if (alt and ri % 2 == 1) else "FFFFFF"
            for ci, cell_text in enumerate(row):
                set_cell_color(cells[ci], fill)
                p = cells[ci].paragraphs[0]
                run = p.add_run(str(cell_text))
                run.font.size = Pt(8.5)
                run.font.name = "STHeiti"
                run._r.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
                run.font.color.rgb = C_DARK
                cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 欄寬
        for ri in range(len(tbl.rows)):
            for ci, w in enumerate(col_widths_cm):
                tbl.rows[ri].cells[ci].width = Cm(w)

        doc.add_paragraph()  # 表後空行

    # ── 封面 ──────────────────────────────────────────────────────────────

    # 封面標題
    p_cover = doc.add_paragraph()
    p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cover.paragraph_format.space_before = Pt(80)
    p_cover.paragraph_format.space_after  = Pt(6)
    set_para_shading(p_cover, NAV_HEX)
    run = p_cover.add_run("\n  投資評級操作手冊\n")
    run.bold = True; run.font.size = Pt(28)
    run.font.name = "STHeiti"
    run._r.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    run.font.color.rgb = C_WHITE

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(p_sub, NAV_HEX)
    run2 = p_sub.add_run("Investment Grade Operating Manual\n")
    run2.font.size = Pt(13); run2.font.name = "STHeiti"
    run2._r.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    run2.font.color.rgb = C_GOLD

    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(p_info, NAV_HEX)
    run3 = p_info.add_run("\nStockGOGO V2  ·  三層決策引擎  ·  Version 2.0  |  2026-04\n ")
    run3.font.size = Pt(9); run3.font.name = "STHeiti"
    run3._r.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    run3.font.color.rgb = C_GRAY

    doc.add_page_break()

    # ── 一、架構概覽 ──────────────────────────────────────────────────────
    add_heading1("一、評級架構概覽")
    add_body("StockGOGO V2 採用「三層過濾 → 評級輸出」架構，每層設有否決門檻，"
             "不通過則直接輸出否決結果，確保分析邏輯清晰、各層因子不互相干擾。")

    add_heading3("流程架構")
    add_table(
        ["", "層級", "評分維度", "核心問題", "輸出"],
        [
            ["", "Layer 1", "方向分", "均線排列是否偏多？", "< 40 → SKIP"],
            ["", "Layer 2", "位置分", "現在位置是否合理？", "< 40 → WAIT"],
            ["", "Layer 3", "時機分", "今天是否有進場觸發？", "→ A / B / C / X"],
            ["", "賣訊", "優先", "三盤跌破 / 形態頭部 / 籌碼惡化", "→ SELL / HOLD"],
        ],
        [0.5, 1.5, 1.5, 7.0, 4.5]
    )

    add_heading3("評級總覽")
    add_table(
        ["評級", "動作代碼", "觸發層", "分數區間", "動作摘要"],
        [
            ["A 主攻",   "STRONG_BUY", "Layer 3",   "70–100", "立即進場，50–80% 部位"],
            ["B 追蹤",   "BUY",        "Layer 3",   "55–69",  "輕倉試單，等確認再加碼"],
            ["C 觀察",   "HOLD",       "Layer 3",   "40–54",  "列入名單，目前不開倉"],
            ["X 無訊號", "HOLD",       "Layer 3",   "40–54",  "同 C 級，不操作"],
            ["WAIT",     "WAIT",       "Layer 2 否決", "30–49", "趨勢對但位置過熱，等拉回"],
            ["SKIP",     "SKIP",       "Layer 1 否決", "0–39",  "趨勢空頭，完全跳過"],
            ["SELL 防守","SELL",       "賣訊優先",  "—",     "三盤跌破，次日開盤全出"],
            ["SELL 停利","HOLD/SELL",  "賣訊優先",  "—",     "過熱賣訊，分批獲利了結"],
            ["SELL 反轉","SELL",       "賣訊優先",  "—",     "頭部確立或法人連賣 7 天"],
        ],
        [2.0, 2.8, 2.5, 1.8, 6.4]
    )
    doc.add_page_break()

    # ── 二、買進評級 ──────────────────────────────────────────────────────
    add_heading1("二、買進評級詳細說明")

    add_heading2("A 級  主攻（立即進場）")
    add_body("方向分通過 + 位置分通過 + 今天有明確的量價觸發訊號。三層全部到位，是最強力的進場訊號。", "定義：")
    add_heading3("觸發條件（任一即可）")
    for t in ["三盤突破 + 帶量確認（收盤突破近三日高點，成交量明顯放大）",
              "VP05 帶量突破訊號（量價分析系統確認）",
              "底部形態（W底 / 頭肩底）突破頸線，且為剛確立狀態"]:
        add_bullet(t)
    add_heading3("操作動作")
    add_table(["項目", "建議"],
              [["進場時機", "當日或次日開盤，不等待"],
               ["部位大小", "50–80% 標準部位"],
               ["停損設定", "進場前低點，或 MA20 下方 2–3%"],
               ["停利目標", "以 RR ≥ 2 為基準，參考形態測幅"],
               ["注意事項", "當日已漲 >5%，降為 B 級，等回檔再進"]],
              [3.5, 12.0])
    add_note("A 級不是「感覺不錯」，而是量價到位。方向對但沒有觸發訊號，最多只是 B 級。")

    add_heading2("B 級  追蹤（等待確認）")
    add_body("方向分通過 + 位置分通過 + 訊號形成但尚未完整確認。值得買，但今天不是最佳時機。", "定義：")
    add_heading3("觸發條件（任一即可）")
    for t in ["三盤突破但量能不足（無帶量確認）",
              "超跌反彈左側訊號（RSI 超賣區出現止跌量）",
              "底部形態形成中，且距離頸線 < 3%"]:
        add_bullet(t)
    add_heading3("操作動作")
    add_table(["項目", "建議"],
              [["進場時機", "等待：① 帶量突破確認 ② 次日跳空開高 ③ 量縮回踩不破支撐"],
               ["部位大小", "30–50% 輕倉試單，升格 A 後再加碼"],
               ["停損設定", "比 A 級稍寬，設在近期低點或 MA60"],
               ["停利目標", "RR ≥ 1.5，目標相對保守"],
               ["加碼條件", "出現 A 級訊號後，才加碼到滿倉"]],
              [3.5, 12.0])
    add_note("B 級不是「馬上買」，是「值得盯著，等訊號」。停損設太緊容易被洗，建議略寬。")

    add_heading2("C 級  觀察 ／ X 無訊號（不操作）")
    add_body("方向分與位置分均通過，但時機層沒有具體觸發條件。C 為多頭環境但無催化，X 為連多頭環境都不明確。", "定義：")
    add_heading3("操作動作")
    add_table(["項目", "說明"],
              [["現在做什麼", "加入自選股，設定價格提醒"],
               ["升格條件",  "帶量突破 → 升 A；左側訊號或接近頸線 → 升 B"],
               ["已持有倉位","無賣訊則繼續持有，但不加碼"],
               ["X vs C",    "X 比 C 更不確定，升格機率較低，優先序較後"]],
              [3.5, 12.0])
    doc.add_page_break()

    # ── 三、過濾評級 ──────────────────────────────────────────────────────
    add_heading1("三、過濾評級詳細說明")

    add_heading2("WAIT  等待拉回（Layer 2 否決）")
    add_body("方向分通過（趨勢向上），但位置分 < 40（位置過熱或 RR 不足）。系統幫你踩煞車，避免追高。", "定義：")
    add_heading3("常見觸發情況")
    for t in ["乖離率 > +10%（距 MA20 超漲 10%）", "RSI > 80（技術面嚴重過熱）",
              "風險報酬比 < 1.5（上方空間不足）", "乖離率 > +15%（天花板啟動，位置分硬封頂 ≤ 20）"]:
        add_bullet(t)
    add_heading3("操作動作")
    add_table(["項目", "說明"],
              [["現在做什麼", "絕對不追，耐心等待"],
               ["等什麼",    "乖離回歸、RSI 降溫、或出現回檔量縮"],
               ["升格條件",  "乖離回到 -5%~+3%，RSI 回到 40–65，升為 B 或 A"],
               ["最常犯錯誤","「趨勢很好應該繼續漲」—— 此理由在 WAIT 時完全無效"]],
              [3.5, 12.0])
    add_note("WAIT 讓人痛苦，但追高後套牢的損失，遠大於沒買到的機會成本。")

    add_heading2("SKIP  跳過（Layer 1 否決）")
    add_body("均線方向分 < 40，趨勢反對你，任何買進都是逆勢操作。", "定義：")
    add_heading3("常見觸發情況")
    for t in ["均線完全空頭排列（MA20 < MA60 < MA120 < MA240）",
              "現價在年線以下", "4 條均線中，多頭方向 ≤ 1 條"]:
        add_bullet(t)
    add_heading3("操作動作")
    add_table(["項目", "說明"],
              [["現在做什麼", "完全不看這檔，從關注名單移除"],
               ["重新關注條件","方向分升到 40 以上（至少 2/4 均線轉多頭）"],
               ["已持有倉位", "SKIP 本身不是賣訊，需同時出現賣訊評級才出場"],
               ["最常犯錯誤", "「已經跌很多應該反彈」—— SKIP 時此理由無效"]],
              [3.5, 12.0])

    add_heading3("SKIP vs. WAIT 差異比較")
    add_table(["比較項目", "SKIP", "WAIT"],
              [["問題所在", "趨勢空頭或混亂",    "位置過熱，追高風險"],
               ["趨勢狀態", "空頭確認",           "多頭確認"],
               ["解決辦法", "等趨勢轉向（數週）", "等位置拉回（數天）"],
               ["緊急程度", "完全退場觀察",       "耐心等待位置"]],
              [3.5, 6.5, 5.5])
    doc.add_page_break()

    # ── 四、賣出評級 ──────────────────────────────────────────────────────
    add_heading1("四、賣出評級詳細說明")
    add_body("賣訊優先於所有買訊。三種賣出類型各有明確觸發條件與執行節奏，不混用、不猜測。")

    add_heading2("SELL 防守型  立即出場")
    add_body("已有倉位，現在繼續持有的風險已超過潛在獲利。最高優先級，須立即執行。", "定義：")
    add_heading3("觸發條件（任一即可）")
    for t in ["三盤跌破（收盤跌破近三日低點，趨勢轉空）",
              "跌破 MA20 超過 2%（收盤在 MA20 × 0.98 以下）"]:
        add_bullet(t)
    add_heading3("操作動作")
    add_table(["項目", "建議"],
              [["執行速度", "收到訊號後，次日開盤確認後即出"],
               ["部位處理", "全部出場，不留半倉「賭反彈」"],
               ["出場方式", "掛市價或略低於市價的限價，確保成交"],
               ["之後動作", "等系統重新評為 B 或 A，再考慮重新進場"],
               ["最常犯錯誤","「跌破了但明天會反彈」—— 嚴格執行停損是必要紀律"]],
              [3.5, 12.0])
    add_note("防守型賣訊出現時，繼續持有的期望值已為負。停損後漲回的痛苦遠小於持續虧損的損失。")

    add_heading2("SELL 停利型  獲利了結")
    add_body("已有獲利，技術面顯示上漲動能衰竭。主動管理倉位，不是緊急出場。", "定義：")
    add_heading3("觸發條件（任一即可）")
    for t in ["過熱賣訊：乖離率 > +15%，且出現長上影線或爆量不漲",
              "移動停利觸發：從最高點回落 5%，或跌破 5 日線"]:
        add_bullet(t)
    add_heading3("操作動作")
    add_table(["項目", "建議"],
              [["執行方式", "不必急，可分批出場（先減 50%，剩餘跟移動停利）"],
               ["是否全出", "視個股強弱，強勢股可留 20–30% 看後續"],
               ["出場標準", "以計畫中的目標價為基準，達標就出"],
               ["最常犯錯誤","「還會繼續漲」而遲遲不出，最後停利變停損"]],
              [3.5, 12.0])

    add_heading2("SELL 反轉型  趨勢反轉出場")
    add_body("整體趨勢已轉向，不只是短線回調，而是結構性改變。代表多頭週期可能結束。", "定義：")
    add_heading3("觸發條件（任一即可）")
    for t in ["頭部形態確立（M頭 / 頭肩頂跌破頸線確認）",
              "法人連續賣超 ≥ 7 天（籌碼嚴重惡化）"]:
        add_bullet(t)
    add_heading3("操作動作")
    add_table(["項目", "建議"],
              [["執行方式", "視情況分批或全部出場"],
               ["頭部形態觸發", "跌破頸線當天或次日，盡速出場"],
               ["籌碼惡化觸發", "觀察 2–3 天，趨勢未改善則出場"],
               ["重新進場條件", "等評級重新回到 B 或 A，且趨勢明確轉多"],
               ["最常犯錯誤", "「基本面沒變壞」—— 技術面反轉先跑，等基本面確認往往太慢"]],
              [3.5, 12.0])

    add_heading3("三種賣訊比較")
    add_table(["類型", "動作代碼", "觸發條件", "緊急程度", "執行方式"],
              [["防守型", "SELL（urgent）", "三盤跌破 / 跌破 MA20", "高",   "全部出場"],
               ["停利型", "HOLD / SELL",   "過熱賣訊 / 移動停利",  "中",   "分批出場，可留倉"],
               ["反轉型", "SELL（urgent）", "頭部確立 / 法人連賣",  "中高", "盡速出場"]],
              [2.0, 2.8, 5.5, 1.8, 3.4])
    doc.add_page_break()

    # ── 五、倉位管理 ──────────────────────────────────────────────────────
    add_heading1("五、倉位管理與籌碼輔助")

    add_heading3("評級對應的倉位建議")
    add_table(["評級", "最大持倉", "說明"],
              [["A 主攻",  "50–80%", "三層全過 + 強觸發，積極進場"],
               ["B 追蹤",  "30–50%", "等確認再加碼，先輕倉"],
               ["C / X",   "0%",     "不開倉，列入追蹤"],
               ["WAIT",    "0%",     "不追高，等位置改善"],
               ["SKIP",    "0%",     "完全不看，從名單移除"],
               ["賣訊出現","0%",     "依類型執行出場"]],
              [2.5, 2.2, 10.8])

    add_heading3("加碼邏輯（唯一合理路徑）")
    add_table(["步驟", "條件", "動作"],
              [["步驟 1", "B 級出現",   "輕倉進場 30–50%"],
               ["步驟 2", "升格為 A 級","加碼至 50–80%"],
               ["步驟 3", "獲利達 +10%","啟動移動停利追蹤"],
               ["步驟 4", "停利型賣訊", "分批出場"]],
              [1.8, 4.5, 9.2])
    add_note("評級下降（A→B→C）時，不加碼，考慮減碼。出現任何 SELL 訊號時，只考慮出場。")

    add_heading3("籌碼的輔助作用")
    add_table(["籌碼狀態", "排序加成", "實際效果"],
              [["法人連買 ≥ 5 天", "+2", "同評級中優先選擇"],
               ["法人連買 3–4 天", "+1", "次優先"],
               ["中性",           " 0", "正常排序"],
               ["法人連賣 3–4 天", "−1", "同評級中後選"],
               ["法人連賣 ≥ 5 天", "−2", "B/C 級自動降一級"]],
              [5.5, 2.0, 8.0])
    add_note("籌碼不決定買賣，只做排序與過濾。再好的籌碼，趨勢空頭時仍為 SKIP。")
    doc.add_page_break()

    # ── 六、情境對照 ──────────────────────────────────────────────────────
    add_heading1("六、常見情境對照 ／ 評級升降格含義")

    add_heading3("常見情境對照")
    add_table(["情境", "評級", "動作"],
              [["台積電均線多頭，三盤突破帶量",      "A 主攻",   "進場 50–80%"],
               ["某股趨勢好，但 RSI=82、乖離+12%",  "WAIT",     "不追，等拉回"],
               ["某股剛從低點反彈，W底接近頸線",     "B 追蹤",   "輕倉試單，等突破"],
               ["某股年線以下，均線死亡排列",        "SKIP",     "完全不看"],
               ["持有中股票出現三盤跌破",            "SELL 防守","次日開盤出場"],
               ["持有中股票漲15%後爆量長上影線",     "SELL 停利","分批出場"],
               ["持有中股票出現 M 頭跌破頸線",       "SELL 反轉","盡速出場"],
               ["持倉評級從 A 降到 C",              "C 觀察",   "不加碼，持有觀察"],
               ["同時有 3 檔 A 級，資金有限",        "A 主攻",   "按籌碼 ranking_boost 排序"],
               ["評級 B + 法人連賣 6 天",           "C 觀察",   "自動降級，降低優先序"]],
              [8.5, 2.5, 4.5])

    add_heading3("評級升降格含義")
    add_table(["方向", "路徑", "含義"],
              [["升格", "SKIP → WAIT",    "趨勢開始轉多，但尚不值得追"],
               ["升格", "WAIT → B/C",     "位置回到合理範圍，進入觀察"],
               ["升格", "C/B → A",        "量價確認，可以動手"],
               ["降格", "A → B",          "觸發訊號消失，可能假突破，考慮減倉"],
               ["降格", "B → C",          "環境轉差或訊號失效，持有但不加碼"],
               ["降格", "C → WAIT/SKIP",  "趨勢或位置惡化，已持有者考慮減碼"]],
              [1.8, 3.8, 9.9])
    add_note("降格不等於馬上賣出。降格是警示，賣出需同時出現 SELL 賣訊。"
             "沒有賣訊的降格代表風險上升，但尚未到出場時機。")

    # 版權
    doc.add_paragraph()
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(p_end, NAV_HEX)
    r = p_end.add_run(
        "本手冊為 StockGOGO V2 三層決策引擎配套說明文件  ·  "
        "Version 2.0  ·  2026-04  ·  機密文件，僅供內部使用"
    )
    r.font.size = Pt(8); r.font.name = "STHeiti"
    r._r.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    r.font.color.rgb = C_GRAY

    doc.save(output_path)
    print(f"Word 已輸出：{output_path}")


# ══════════════════════════════════════════════════════════════════════════════
#  主程式
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    base = os.path.dirname(os.path.abspath(__file__))
    pdf_path  = os.path.join(base, "投資評級操作手冊.pdf")
    word_path = os.path.join(base, "投資評級操作手冊.docx")

    print("正在產生 PDF...")
    build_pdf(pdf_path)

    print("正在產生 Word...")
    build_word(word_path)

    print("\n完成！")
    print(f"  PDF  → {pdf_path}")
    print(f"  Word → {word_path}")
